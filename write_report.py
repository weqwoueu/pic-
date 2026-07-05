from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document(r'E:\作业\专业实践\研究生专业实践报告模板.docx')

# Fix cover page fields
cover_fields = {
    '业：': '业：  等离子体物理 / 计算物理',
    '称：': '称：  电子速度分布对激光等离子体膨胀\n                         影响的全动理学PIC模拟研究',
    '报告人': '报告人（学号）：  dg001947',
    '期：': '期：  2026年7月',
}

for p in doc.paragraphs:
    for key, val in cover_fields.items():
        if key in p.text:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = val
            break

def add_heading_styled(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(size)
    run.bold = True

def sec(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(14)
    run.bold = True

def sub(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12)
    run.bold = True

def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

doc.add_page_break()

# ==================== CONTENT ====================

sec(doc, '一、实践课题概述')

sub(doc, '1.1 课题名称')
body(doc, '电子速度分布对激光等离子体膨胀影响的全动理学PIC模拟研究。')

sub(doc, '1.2 课题来源')
body(doc, '本课题来源于实验室课题组承担的激光等离子体相互作用数值模拟研究项目。课题以发表在等离子体物理领域的学术论文"Electron distribution effects on laser-plasma expansion with fully kinetic simulation"为核心参考，旨在通过全动理学粒子网格（Particle-in-Cell, PIC）数值模拟方法，系统研究不同电子速度分布函数对无碰撞等离子体向真空膨胀过程的影响。')

sub(doc, '1.3 实践目标')
body(doc, '本次专业实践的主要目标包括：第一，深入理解全动理学PIC模拟的基本原理与数值方法，掌握浸入式有限元（Immersed Finite Element, IFE）方法在复杂边界条件下求解泊松方程的技术；第二，参与课题组现有的IFE-PIC代码的维护与功能扩展，为代码添加全局能量守恒与粒子数守恒诊断模块；第三，按照指导老师提出的数值验证方案，系统开展网格分辨率、时间步长、粒子数以及随机种子等维度的收敛性测试；第四，为后续完整复现论文中的十四张核心对比图并补充收敛性验证图件奠定基础。')

sec(doc, '二、研究背景与意义')

sub(doc, '2.1 物理问题背景')
body(doc, '等离子体向真空的膨胀是众多物理领域中的基本过程，涉及离子加速、空间等离子体物理、离子注入以及极紫外（EUV）光刻光源等关键应用。当高功率激光辐照固体靶材时，会产生高温高密度的激光等离子体（Laser-Produced Plasma, LPP）。等离子体在向真空膨胀过程中，电子与离子之间通过自洽电场耦合，电子的热力学行为直接决定了离子加速的能量和膨胀前沿的传播速度。')
body(doc, '传统研究通常假设电子遵循经典的麦克斯韦-玻尔兹曼（Maxwell-Boltzmann）速度分布。然而，近年来高功率激光-固体相互作用的实验表明，高能电子群体可能显著偏离热平衡态，呈现出非麦克斯韦分布特征，尤其是具有高能尾部的Kappa分布以及具有速度截断的Polytropic（多方）分布。这些非平衡电子分布会导致等离子体热力学行为发生质的变化，直接影响离子加速效率与膨胀动力学。')

sub(doc, '2.2 三种电子速度分布')
body(doc, '本研究聚焦于三种典型的电子速度分布函数：')
body(doc, '麦克斯韦分布（Maxwellian Distribution）是最经典的平衡态分布，其对应的多方指数γe = 1，在膨胀过程中维持等温行为——电子温度在未扰动区域保持恒定。麦克斯韦分布作为基准参照，用于与其他非平衡分布进行对比。')
body(doc, 'Kappa分布广泛用于描述空间等离子体（如太阳风、日冕、行星磁层）中包含高能粒子尾部的系统。Kappa分布的谱指数κ表征偏离麦克斯韦的程度，较小的κ值对应更强的非热力学特征。当κ → ∞时，Kappa分布趋近于麦克斯韦分布。Kappa分布的多方指数γe < 1，这意味着在等离子体膨胀过程中电子温度反而升高，呈现所谓的"反热力学"（inverted thermodynamics）行为——密度降低但温度上升，这完全违背了常规绝热冷却的直觉。')
body(doc, 'Polytropic（多方）分布由Kiefer等人提出，用多方指数γe表征偏离等温的程度。γe > 1时电子在膨胀中正常冷却，分布函数具有速度截断特征，即不存在超过某一最大速度的粒子。当γe = 3时，该分布退化为阶跃函数形式，此时能量的守恒性在理论上得到严格保证，但数值上极为敏感。')

sub(doc, '2.3 研究的科学意义')
body(doc, '三种分布在热力学上表现出质的差异——Kappa分布的反常加热（γe < 1）、麦克斯韦分布的等温行为（γe = 1）、Polytropic分布的正常冷却（γe > 1）——构成了一个统一的热力学框架。然而，大多数前期研究基于流体力学模型并忽略了有限离子质量带来的电子惯性效应。Denavit等人的全动理学模拟先驱工作证明了离子质量会显著修正非等温等离子体的温度演化。本课题正是立足于全动理学PIC模拟，系统考察有限质量比（mi/me = 100、400、900、1600）下不同速度分布的热力学行为，预期揭示∆γe（多方指数偏差）与mi/me之间的标度律关系。')

sec(doc, '三、研究方法与技术路线')

sub(doc, '3.1 全动理学PIC方法')
body(doc, 'PIC方法是等离子体数值模拟中最重要和广泛使用的方法之一。其基本思想是将等离子体视为大量离散的宏粒子（macro-particle），在拉格朗日框架下追踪每个粒子的运动，同时在欧拉框架下的网格上求解电磁场方程。每个时间步内的核心操作包括：将粒子电荷分配到网格节点求解泊松方程获得自洽电场，再将电场插值回粒子位置推动粒子运动。')
body(doc, '本研究使用的IFE-PIC代码采用全动理学（fully kinetic）方法，即电子和离子均被当作离散粒子处理，通过耦合牛顿运动方程与泊松方程来演化等离子体系统。与流体力学模型或混合模型（电子用流体、离子用粒子）相比，全动理学方法能够自然地捕捉电子惯性效应、有限拉莫尔半径效应以及非麦克斯韦速度分布的动力学细节，但也带来了更大的计算开销——电子与离子的质量比越大，为保证数值稳定性所需要的时间步长越苛刻。')

sub(doc, '3.2 浸入式有限元（IFE）方法')
body(doc, '代码采用了浸入式有限元方法在复杂边界上求解泊松方程。IFE方法的核心优势在于它允许网格不完全贴合物体边界（"浸入"边界），通过修改被界面切割的单元的基函数来满足界面跳跃条件。这使得IFE方法在模拟包含复杂电极或靶材构型的等离子体设备时避免了传统贴体网格生成的高昂成本。本代码使用结构化笛卡尔网格（1025×5网格点），模拟域大小为[0, 1024λD0] × [0, 4λD0]，其中λD0为初始德拜长度。IFE网格的分辨率为∆x = ∆y = λD0。电场在离散化网格上求解后，通过补丁多项式重构方法插值到粒子位置用于推动。')

sub(doc, '3.3 归一化方案')
body(doc, '为了应对等离子体物理参数跨越多数量级（如电子质量~10⁻³¹ kg、德拜长度~10⁻⁷ m、等离子体频率~10¹² Hz）带来的数值舍入误差风险，代码采用基于电子尺度的无量纲归一化方案。归一化参考量包括：电子等离子体频率ωpe、德拜长度λD0、电子热速度vth,e0以及参考温度Te0。所有物理量在计算中均以归一化形式存储与处理，仅在输出时按需恢复为有量纲值。')

sub(doc, '3.4 蒙特卡洛碰撞（MCC）模块')
body(doc, '代码集成了蒙特卡洛碰撞模块（MCC_jw），用于模拟电子/离子与中性气体原子之间的碰撞过程。尽管本课题研究的核心算例为无碰撞等离子体膨胀（背景气体压强设为零），MCC模块提供的粒子初始化框架（包括速度分布采样、粒子束管理、边界处理等）是代码的重要组成部分。碰撞截面数据库采用LXCat及Pegasus格式，涵盖He、Ar、N₂、CO₂等多种气体。')

sec(doc, '四、实践工作内容')

sub(doc, '4.1 代码学习与结构梳理')
body(doc, '在实践初期，我对课题组已有的IFE-PIC代码进行了全面的梳理与学习。代码主体使用Fortran语言编写，采用Intel Fortran编译器（ifort 2021.5.0）编译，构建系统基于CMake（版本 ≥ 3.10）。代码结构分为以下几个主要部分：')
body(doc, 'code/PIC/——主PIC算法，包括主时间循环（Main_IFE_Test_2.f90）、粒子推动、电荷分配、边界条件设置等；code/IFE-Rectangular/与code/IFE-Solver/——浸入式有限元的基函数计算、刚度矩阵装配、迭代求解器（预共轭梯度法PCG）；code/SIDG/——选择浸入式间断伽辽金（Selective Immersed Discontinuous Galerkin）扩展，实现自适应网格细化与场重构；code/PPR/——补丁多项式重构方法，用于从网格节点场值插值获得粒子位置的电场值；MCC_jw/——蒙特卡洛碰撞模块，包含粒子类型定义、碰撞截面计算、反应通道处理、边界处理以及多种诊断功能；code/In-Output/——输入参数读取与Tecplot格式场量输出。根级诊断文件OUTPUT_velocity.f90和Output_Energy.f90分别负责漂移/热速度计算和能量分布函数计算。')
body(doc, '在代码梳理过程中，我重点分析了主时间循环的执行流程、粒子逃逸与注入的边界处理逻辑、电场能量的体积积分方法、以及不同速度分布（Maxwellian、Kappa、Polytropic）的初始化采样算法。这些工作为后续的功能扩展打下了基础。')

sub(doc, '4.2 全局能量与粒子数守恒诊断模块开发')
body(doc, '按照指导老师的工作建议，验证全动理学PIC模拟结果的可信度需要在每个算例中输出完整的能量守恒与粒子数守恒诊断信息。我独立完成了全局诊断模块的开发，包括两个核心Fortran源文件：')
body(doc, 'DiagCounters.f90——独立的共享计数器模块。该模块定义了整个程序中能量损失/注入和粒子损失/注入的累积计数器（E_lost_cum、E_inj_cum、Ne_lost_cum、Ni_lost_cum、Ne_inj_cum、Ni_inj_cum），并提供重置和累加的子程序接口（ResetDiagCounters、AddDiagLost、AddDiagInjected）。该模块独立于所有其他模块，不产生循环依赖，遵循了软件工程中模块解耦的设计原则。')
body(doc, 'GlobalDiagnostics.f90——全局诊断主模块。该模块在每个时间步完成以下计算：遍历所有粒子计算电子动能Ke和离子动能Ki（每个宏粒子的动能乘以统计权重后求和，单位为焦耳）；对归一化电场分量efx、efy进行去归一化并做体积积分（利用Efield_ref = Phi_ref / L_ref和Cell_Volume_zwz），计算电场能量Efield = ½ε₀ ∫ (Ex² + Ey²) dV；读取DiagCounters模块中的累积损失/注入值；计算总能量E_dom = Ke + Ki + Efield；输出相对能量守恒误差δE(t) = [E_dom(t) + E_lost(t) - E_inj(t) - E_dom(0)] / E_dom(0)。')
body(doc, '诊断输出文件为OUTPUT/global_diagnostics.csv，格式严格遵循指导老师规定的十三列数据：归一化时间t、域内电子数Ne_domain、域内离子数Ni_domain、累积损失电子数Ne_lost、累积损失离子数Ni_lost、累积注入电子数Ne_injected、累积注入离子数Ni_injected、电子动能Ke、离子动能Ki、电场能Efield、累积注入能量Einjected、累积损失能量Elost、能量守恒误差Ebalance_error。该格式设计可直接导入Python（pandas.read_csv）或MATLAB进行后处理和可视化。')

sub(doc, '4.3 粒子损失跟踪的代码集成')
body(doc, '为准确追踪粒子从计算域逃逸所携带的能量，我在MCCInterface.f90的MoveAll子程序中添加了粒子删除前的能量累加逻辑。当粒子位置标记为X ≤ dxmin（所有逃逸粒子的统一标记）时，在调用DelOne删除之前，计算该粒子的动能（½ m v² × VFactor²）并按物种分类（电子为SpeciesIndex = 0，离子为SpeciesIndex ≥ 1）调用AddDiagLost累加到DiagCounters模块的相应计数器。该方案通过在粒子删除点进行能量采样，避免了事后统计的不确定性。这一修改不改变原有的粒子删除逻辑，仅在删除前增加了一次能量采样，对性能的影响微乎其微（单个粒子的能量计算仅涉及几次浮点乘法运算）。')

sub(doc, '4.4 代码部署与编译环境搭建')
body(doc, '在代码开发完成后，我负责在电光云HPC集群（cloud.dghpc.com）上完成了代码的部署与编译。具体工作包括：配置SSH密钥认证实现免密登录；使用Git进行代码版本管理，从GitHub仓库拉取最新代码；在Rocky Linux 8.10环境下加载Intel Fortran编译器（ifort 2021.5.0，通过module load intel加载）和CMake 3.24.2；运行setup_maxwell_mi400_case.sh脚本生成麦克斯韦分布mi/me=400算例的输入文件；通过CMake构建系统完成代码编译（共编译约100个Fortran源文件，链接生成单一可执行文件1DPIC）。')
body(doc, '编译过程顺利通过，无编译错误和警告。运行时遇到段错误（SIGSEGV），经traceback调试（使用-O0 -g -traceback编译选项）定位在粒子初始化子程序AllInitializationParticleBundleIFE的PosInit调用（ParticleOne.f90:86）。初步判断为setup脚本的perl补丁在git clone得到的代码版本上不完全兼容所致。该问题与全局诊断模块无关（注释掉诊断模块后崩溃位置和堆栈完全相同），目前正在通过恢复此前本地scp传输的可运行版本并结合诊断模块增量修改来解决。')

sub(doc, '4.5 算例参数准备')
body(doc, '根据论文及指导老师的技术要求，基准算例参数设置如下：等离子体初始占据[0, 128λD0] × [0, 4λD0]的矩形区域，嵌入[0, 1024λD0] × [0, 4λD0]的全模拟域；电子温度Te0 = 1 eV，离子温度Ti0 = 0.01 eV（温度比Te0/Ti0 = 100）；离子带单电荷（Z = 1），质量比mi/me = 400；网格间距∆x = ∆y = λD0，时间步长∆t = 0.05 ωpe⁻¹；左边界为镜面反射（绝热条件），右边界为粒子吸收（外流条件），y方向为周期性边界；电子初始速度按麦克斯韦分布采样。')
body(doc, '为后续收敛性测试准备，拟定了以下参数扫描矩阵：网格分辨率Δx = λD0与0.5λD0；时间步长Δt = 0.05 ωpe⁻¹与0.025 ωpe⁻¹；每网格宏粒子数Nppc = 2.0×10⁴、4.0×10⁴、8.0×10⁴；随机种子至少三组（101、202、303）；诊断时刻为ωpit = 30与50。论文级算例采用Nppc = 8.0×10⁴、nt = 20000，每个物种的宏粒子总数达40.96M（电子+离子共计81.92M宏粒子），对计算资源的消耗在电光云集群的承载范围内（380 GB内存/节点）。')

sec(doc, '五、论文理解与需求分析')

sub(doc, '5.1 论文核心图件梳理')
body(doc, '通过对论文全文（12页）和指导老师工作建议（8页）的仔细阅读，我系统梳理了论文涉及的十四张图及其对应的参数组合。论文图件按内容可分为三类：')
body(doc, '第一类为速度分布验证（Fig 1-2），展示理论分布函数与PIC初始化采样的对比，验证三种分布（Maxwellian、Kappa κ=2,6、Polytropic γe=2,3）的初始速度采样是否准确；第二类为分布对比与质量比扫描（Fig 3-9），是论文的核心内容——Fig 3对比三种分布在mi/me=400下的密度和热速度剖面，Fig 4展示Te-ne的log-log拟合曲线（提取γe），Fig 5-9分别对麦克斯韦、Kappa κ=2、κ=6以及Polytropic γe=2、γe=3五种情况在四种质量比（100/400/900/1600）下进行系统扫描，并给出∆γe vs mi/me的非线性拟合结果；第三类为边界条件对比（Fig 10-14），比较等温（漫反射，模拟长脉冲激光的持续加热）与绝热（镜面反射，模拟短脉冲激光结束后的自由膨胀）两种边界条件在各分布下的热力学差异，特别关注γe随时间从ωpit=30到50的动态演化。')

sub(doc, '5.2 收敛性验证需求')
body(doc, '指导老师的工作建议强调，论文中观测到的γe差异必须在数值上被证明是物理效应而非数值误差。需完成的验证维度包括：网格收敛性——证明∆x = λD0足以分辨德拜屏蔽长度和离子前沿的电荷分离结构，若γe在∆x=λD0和0.5λD0之间的差异小于随机种子标准差，可认为网格已收敛；时间步收敛性——证明∆t = 0.05 ωpe⁻¹不引入显式的能量漂移或高能尾部（尤其对Kappa κ=2）的数值弥散；粒子数收敛性——验证γe随Nppc增加而收敛，拟合误差期望按1/√Nppc下降，当4.0×10⁴和8.0×10⁴的结果已趋接近时可认为当前粒子数足够；随机种子统计——至少使用三组不同随机种子（101、202、303），给出γe = γ̄e ± σγe而非单次模拟结果，σγe应小于不同分布间的γe差异值。')

sec(doc, '六、实践收获与体会')

sub(doc, '6.1 专业技能的提升')
body(doc, '通过本次实践，我在以下方面获得了实质性的提升：')
body(doc, '数值模拟能力：从最初的代码阅读者逐步过渡到代码修改者和功能开发者，深入理解了PIC方法的全流程——从粒子初始化（PosInit、VelMaxInit、VelKappaInit、VelPolyInit）、电场求解（IFE/SIDG/PPR三级联用）、粒子推动（显式半步法或隐式预推-后推法）到各类诊断输出（场量、速度、能量、粒子计数）。特别是掌握了大型科研Fortran代码的模块化设计、编译系统配置、Fortran模块循环依赖的排查与解决（通过提取独立的DiagCounters模块打破GlobalDiagnostics与MCCInterface之间的循环依赖），以及调试技巧。')
body(doc, '高性能计算素养：熟悉了HPC集群的使用流程，包括模块系统（module load/avail）、作业调度系统（Slurm的sbatch/squeue/sinfo）、编译器选项的切换、以及SSH密钥管理。电光云集群（cloud.dghpc.com）拥有269个计算节点、每个节点配置48核CPU和380 GB内存、3.5 PB共享存储，是进行大规模等离子体模拟的理想平台。登录节点c2ln05（16核/186 GB内存）用于代码编辑与编译，计算任务通过Slurm提交到comp分区的计算节点执行。')
body(doc, '版本控制与协作：实践了Git在科研代码开发中的标准工作流——从功能分支开发、提交信息规范编写、到GitHub推送（git add、commit、push的完整流程）。这对后续团队协作开发多分布多参数的批处理脚本至关重要。')
body(doc, '理论知识的深化：通过对论文的反复研读，加深了对等离子体动理学理论的理解：三种速度分布的数学形式和物理含义（Maxwellian的指数衰减、Kappa的幂律尾部、Polytropic的速度截断）、多方指数γe作为热力学桥梁的本质、有限质量比导致的电子惯性效应及其与∆γe的标度关系（Maxwellian: β≈-0.5; Kappa κ=2: β≈-0.27; Polytropic γ=2: β≈+0.62反直觉）、等温与绝热边界条件对应的物理场景等。')

sub(doc, '6.2 对科研工作模式的体会')
body(doc, '本次实践让我体会到了科研代码开发的几个特点：第一，科研代码往往带有历史遗留问题（如Main_IFE_Test_2.f90中多处"这些变量未使用，删除？"的开发者注释、大量被注释掉的废弃代码分支），在添加新功能时需要充分理解现有逻辑，避免引入不可预期的副作用；第二，数值模拟的可信度建立在严谨的收敛性验证之上，任何结论都需要附上误差范围和收敛性证据，不能仅凭单次模拟结果下结论——这也是指导老师反复强调"不能只用单次拟合结果支撑细微的∆γe结论"的原因；第三，大型Fortran代码的编译和运行高度依赖于特定的编译器版本和硬件环境，跨平台迁移需要额外的工作量，特别是perl脚本批量修改源文件的做法在不同代码版本间可能存在兼容性风险。')

sub(doc, '6.3 后续工作计划')
body(doc, '基于目前已完成的工作和遇到的问题，后续工作计划如下：')
body(doc, '近期（1-2周）：解决当前段错误问题（恢复scp传输的可运行版本并重新集成诊断模块），跑通基准麦克斯韦mi/me=400验证算例（nt=20000、Nppc=80000），输出完整的global_diagnostics.csv；验证能量守恒误差δE(t) < 1%及粒子数守恒误差δN(t) < 0.1%。')
body(doc, '中期（3-6周）：实现Kappa分布（κ=2）和Polytropic分布（γe=2）的粒子初始化支持（修改MaxKappa标志位和对应的采样子程序）；完成网格分辨率、时间步长、粒子数三个维度的收敛性测试（每个维度涉及3种分布×2-3个参数取值=18-27个算例）；编写γe统一log-log拟合程序（拟合区间10⁻³ ≤ ni/n₀ < 1，排除κ=2的高能尾部极端区）；输出网格收敛对比图、时间步收敛对比图、粒子数收敛趋势图、随机种子误差表。')
body(doc, '远期（2-3月）：复现论文全部十四张图，为每张图补充收敛性验证版本；对Polytropic γe=3（阶跃分布）的特殊行为（∆γe呈现非单调、无量纲依赖）进行专门讨论并归因于速度截断的不连续性；撰写验证结果分析，形成可用于论文正文的数值可信度声明（如"在随机种子统计误差范围内，网格减半和时间步减半不改变γe值"）；将诊断模块、批处理脚本和绘图代码整理为可复用的工具包提交至课题组代码库。')

sec(doc, '七、总结')
body(doc, '本次专业实践以激光等离子体膨胀的全动理学PIC模拟为核心，围绕课题组已有的IFE-PIC代码开展了深入的学习与开发工作。我完成了全局能量与粒子数守恒诊断模块的设计与实现（DiagCounters.f90和GlobalDiagnostics.f90），在电光云HPC集群上成功完成了代码的编译部署，系统梳理了论文中十四张核心图件的参数组合与二十余个验证算例的运行矩阵，并为后续的收敛性测试工作制定了详细的分阶段计划。')
body(doc, '通过本次实践，我不仅提升了Fortran编程、高性能计算和数值模拟的专业技能，更重要的是建立起了对科研代码开发严谨性的认识——任何数值结论都必须建立在充分的收敛性验证和误差分析之上。这一认识将指导我在后续的论文复现和研究工作中，以更高的标准要求自己的每一步数值实验，确保最终产出的科研成果经得起检验。')

# Save
output_path = r'E:\作业\专业实践\研究生专业实践报告_已填写.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
